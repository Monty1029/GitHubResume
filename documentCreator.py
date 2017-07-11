from docx import Document
from docx.shared import Pt

from user import User
from repo import Repo

class DocumentCreator:

    def __init__(self, name=None, user=None, allRepos=None):
        self.name = name
        self.user = user
        self.allRepos = allRepos
    
    def buildDoc(self):
        doc = Document()
        doc.add_heading(self.user.getName(self.name) + " (" + self.name + ")", level = 1)
        p = doc.add_paragraph("")
        p.add_run("Followers: ").bold = True
        p.add_run(self.user.getFollowers(self.name))
        doc.add_heading("Bio", level = 2)
        doc.add_paragraph(self.user.getBio(self.name))
        doc.add_heading("Organizations", level = 2)
        doc.add_paragraph(self.user.getOrganizations(self.name))
        doc.add_heading("Projects", level = 2)
        repoIndex = 0
        for x in self.allRepos:
            p = doc.add_paragraph("")
            p.add_run(x.getRepoName()).bold = True
            p.add_run("\t" + x.getCreationDate())
            p = doc.add_paragraph("")
            p.add_run("Language: ").bold = True
            p.add_run(x.getLang())
            p = doc.add_paragraph("")
            p.add_run("Stars: ").bold = True
            p.add_run(str(x.getStars()) + "\n")
        doc.save('example.docx')